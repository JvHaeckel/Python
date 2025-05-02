# Integração do Python com Word - Como Criar Contratos Automaticamente
# https://www.youtube.com/watch?v=N01MPYL3UVY

from docx import Document

documento = Document(r"C:\Users\joaorocha\Desktop\Py\Learning\Hashtag\Word\Contrato.docx")

name = input ("Digite o nome para o arquivo: ")

# Para não ficar repetitivo criamos um Dicionário
# for paragrafo in documento.paragraphs:
#     paragrafo.text = paragrafo.text.replace('XXXX', name)
#     paragrafo.text = paragrafo.text.replace("YYYY", ben1)
#     paragrafo.text = paragrafo.text.replace("ZZZZ", ben2)
#     paragrafo.text = paragrafo.text.replace("WWWW", ben3)

nome = input("Digite o nome do contratante: ")
item1 = input("Digite seu 1 bem: ")
item2 = input("Digite seu 2 bem: ")
item3 = input("Digite seu 3 bem: ")
dia = input("Digite o dia:  ")
mes = input("Digite o mês: ")
ano = 2025

referencias = {
    "XXXX" : nome,
    "YYYY" : item1,
    "ZZZZ" : item2,
    "WWWW" : item3,
    "DD"   : dia,
    "MM"   : mes,
    "AAAA" : str(ano),
}

# Explica em 13:30
# Percorre todos os parágrafos do documento Word
for paragrafo in documento.paragraphs:
    
    # Dentro de cada parágrafo, pode haver vários pedaços com estilos diferentes (negrito, itálico, etc.)
    for run in paragrafo.runs:
        
        # Para cada código (ex: "XXXX") e seu valor (ex: "João"), O método .items() retorna pares (chave, valor)
        for codigo, valor in referencias.items():
            
            # Se o código estiver nesse pedaço de texto, ele será substituído pelo valor
            if codigo in run.text:
                run.text = run.text.replace(codigo, valor)

documento.save('Contrato de ' + name + " .docx")