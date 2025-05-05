import pandas as pd
from docx import Document
# import tkinter as tk    Embaixo usamos ele mais bonito
import customtkinter as ctk
from tkinter import filedialog

# Configurando as cores padrões
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")  # ou "green", "dark-blue", etc.
janela = ctk.CTk()
janela.configure(fg_color="lightblue")  # muda a cor do fundo


# Função que converte Excel em Word
def converter():
    # Aqui está sendo usada a biblioteca tkinter para abrir uma janela de diálogo para que o usuário escolha um arquivo no computador. 
    # filedialog.askopenfilename():Esse método abre uma janela para o usuário selecionar um arquivo e retorna o caminho completo do 
    # arquivo escolhido. 
    # 
    # filetypes=[("Excel files", "*.xlsx")]:Especifica que a janela de diálogo deve filtrar e mostrar apenas arquivos com extensão 
    # .xlsx, ou seja, apensa arquivos em Excel
    
    arquivo_excel = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    
    if not arquivo_excel:
        return "Selecione arquivo em excel"
    
    # Lê a planilha usando Pandas
    tabela = pd.read_excel(arquivo_excel)

    # Cria o documento Word
    documento = Document()
    
# Acessa o valor inserido pelo usuário em um campo de texto da interface gráfica (feito com o tkinter).

# campo_titulo é um campo de entrada de texto (como uma caixa de texto) criado na interface do programa.
#.get() - método que obtém o texto  digitado pelo usuário nesse campo, e o armazena na variável titulo.
    titulo = titulo1.get()
    
# Adiciona um título ao documento Word.

# add_heading() é um método da biblioteca python-docx que permite adicionar títulos de diferentes níveis (semelhante aos 
# títulos em um processador de texto como o Word).

# titulo - é o texto que será usado como título, que foi obtido na linha anterior a partir do campo de entrada.
# level=1- Define o nível do título. No Word, o nível 1 corresponde ao título mais importante.
    documento.add_heading(titulo, level=1)

    for _, linha in tabela.iterrows():
        texto = ""
        for coluna in tabela.columns:
            texto += f"{coluna}: {linha[coluna]}  "
        documento.add_paragraph(texto)
        documento.add_paragraph("")  # Linha em branco entre os itens

    # Salvar com nome digitado
    nome_arquivo = titulo2.get()
    if nome_arquivo.strip() == "":
        nome_arquivo = "documento_word"
        
    documento.save(f"{nome_arquivo}.docx")
    

#  *****   Campos na janela  *****

# Criando janela
janela = ctk.CTk()                          # Esse T sempre maiúsculo
janela.geometry("300x225")                # Vai dar o tamanho da caixa
janela.title("Conversor Excel → Word ")   # Título da caixa

# Texto do  campo 1
titulo1 = ctk.CTkLabel(janela, text = "Digite o Título do documento:")
titulo1.pack(pady=(5,5))

# Campo 1
campo1 = ctk.CTkEntry(janela, width=250 , placeholder_text="Digite aqui ")
campo1.pack(pady=(0,5))

# Texto do Campo 2
titulo2 = ctk.CTkLabel(janela , text="Digite o nome do arquivo Word:")
titulo2.pack(pady=(15,5))

# Campo 2
campo2 = ctk.CTkEntry(janela, width=250, placeholder_text="Digite aqui ")
campo2.pack(pady=(0,15))

# Botão para converter
botao = ctk.CTkButton(janela, text="Converter", width=80, height=25, command=converter) # command - chama a função acima
botao.pack(pady=(15,7.5))

# Ele faz com que a janela fique aberta, escutando e respondendo a eventos do usuário
janela.mainloop()