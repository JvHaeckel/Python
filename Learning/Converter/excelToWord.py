# Importações

import pandas as pd
from docx import Document

caminho = (r"C:\Users\joaorocha\Desktop\Py\Learning\Converter\Informações.xlsx")

# Ler e carrega os dados do excel
tabela = pd.read_excel(caminho)

documento = Document()

titulo = input("Qual título do documento? ")
documento.add_heading(titulo)

# Itera pelas linhas da planilha
for index, linha in tabela.iterrows():
    texto = ""
    for coluna in tabela.columns:
        texto += f"{coluna}: {linha[coluna]}  |  "
    documento.add_paragraph(texto.strip(" | "))

# Salva o documento Word
documento.save("planilha_convertida.docx")



